# from docx import Document

# def extract_text_from_docx(docx_file):
#     try:
#         # Open the Word document
#         doc = Document(docx_file)
        
#         # Extract text from paragraphs
#         text_content = []
#         for paragraph in doc.paragraphs:
#             text_content.append(paragraph.text)
        
#         # Join the paragraphs into a single string
#         return "\n".join(text_content)
    
#     except Exception as e:
#         print(f"An error occurred: {e}")
#         return ""


# # Example usage
# docx_file_path = r"C:/Users/rahul/OneDrive/Desktop/shivam/DDM Assignment-2 (1).docx"  
# extracted_text = extract_text_from_docx(docx_file_path)

# print("Extracted Text:")
# print(extracted_text)

# # Save the extracted text to a .txt file (optional)
# with open("extracted_text.txt", "w", encoding="utf-8") as file:
#     file.write(extracted_text)


import os
from docx import Document
from PIL import Image
import pytesseract
from io import BytesIO

# Configure Tesseract path if needed
pytesseract.pytesseract.tesseract_cmd = r"C:/Program Files/Tesseract-OCR/tesseract.exe"

def extract_images_from_docx(docx_file, output_folder):
    """
    Extracts images from a .docx file and saves them to the output folder.
    Returns a list of image file paths.
    """
    doc = Document(docx_file)
    image_paths = []
    
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:  # Check for image relationships
            image_data = rel.target_part.blob  # Get the image binary
            image_path = os.path.join(output_folder, f"image_{len(image_paths) + 1}.png")
            with open(image_path, "wb") as img_file:
                img_file.write(image_data)
            image_paths.append(image_path)
    
    return image_paths

def extract_text_from_images(image_paths):
    """
    Extracts text from a list of image file paths using Tesseract OCR.
    Returns the combined extracted text.
    """
    extracted_text = []
    
    for image_path in image_paths:
        try:
            text = pytesseract.image_to_string(Image.open(image_path))
            extracted_text.append(text.strip())
        except Exception as e:
            print(f"Error processing {image_path}: {e}")
    
    return "\n".join(extracted_text)

# Main function
def extract_text_from_docx_images(docx_file, output_folder):
    # Create the output folder if it doesn't exist
    os.makedirs(output_folder, exist_ok=True)

    # Extract images from the .docx file
    image_paths = extract_images_from_docx(docx_file, output_folder)
    print(f"Extracted {len(image_paths)} images.")

    # Perform OCR on the extracted images
    extracted_text = extract_text_from_images(image_paths)
    
    # Clean up images after processing (optional)
    for image_path in image_paths:
        os.remove(image_path)
    
    return extracted_text

# Example usage
docx_file_path = r"C:/Users/rahul/OneDrive/Desktop/shivam/binder2.docx"  
output_dir = "output_images"

extracted_text = extract_text_from_docx_images(docx_file_path, output_dir)

print("Extracted Text from Images:")
print(extracted_text)

# Save the extracted text to a file if needed
with open("extracted_text_from_images.txt", "w", encoding="utf-8") as text_file:
    text_file.write(extracted_text)

